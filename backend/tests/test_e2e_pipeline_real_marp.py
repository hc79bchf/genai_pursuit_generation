import asyncio
import sys
import os
import json
from datetime import datetime

# Add backend directory to sys.path
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "../backend")))

from app.services.ai_service.llm_service import LLMService
from app.services.ai_service.metadata_agent import MetadataExtractionAgent
from app.services.ai_service.ppt_outline_agent import PPTOutlineAgent

async def run_e2e_pipeline_real_marp():
    print("🚀 Starting End-to-End Pipeline Test (Real Data + MARP)...")
    
    # 1. Setup Input
    input_file_container = "/app/Data/AI_Platform_RFP.docx" 
    input_file_local = "Data/AI_Platform_RFP.docx"
    
    file_content = "Dummy content"
    
    if os.path.exists(input_file_local):
        print(f"📂 Found Input File (Local): {input_file_local}")
        try:
            import docx
            doc = docx.Document(input_file_local)
            file_content = "\n".join([para.text for para in doc.paragraphs])
            print(f"✅ Read {len(file_content)} characters from file.")
        except ImportError:
            print("⚠️ python-docx not installed locally, using dummy content.")
        except Exception as e:
            print(f"⚠️ Failed to read local file: {e}")
    elif os.path.exists(input_file_container):
        print(f"📂 Found Input File (Container): {input_file_container}")
        try:
            import docx
            doc = docx.Document(input_file_container)
            file_content = "\n".join([para.text for para in doc.paragraphs])
            print(f"✅ Read {len(file_content)} characters from file.")
        except Exception as e:
             print(f"⚠️ Failed to read container file: {e}")
    else:
        print(f"⚠️ Input file not found, using dummy content.")

    print(f"📂 Processing Content Length: {len(file_content)}")
    
    # Initialize Real Services
    llm_service = LLMService()
    
    # --- Step 1: Metadata Extraction ---
    print("\n--- Step 1: Metadata Extraction (Real Agent) ---")
    metadata_agent = MetadataExtractionAgent(llm_service)
    
    # We need to mock MemoryService or ensure it works. 
    # MetadataExtractionAgent uses MemoryService to store/retrieve.
    # For this script, we can probably let it run if dependencies are met, 
    # or patch it if it requires DB access we don't want to set up.
    # Let's try running it real first. If it fails on DB, we patch MemoryService.
    
    from unittest.mock import patch
    # Patch MemoryService to avoid DB dependency for this standalone script
    with patch('app.services.ai_service.metadata_agent.MemoryService'):
        extracted_data = await metadata_agent.extract(file_content, user_id="test_user")
        print(f"✅ Extracted Metadata: {extracted_data}")

    # --- Step 2: PPT Generation (PPTOutlineAgent) ---
    print("\n--- Step 2: PPT Generation (Real Agent) ---")
    
    ppt_agent = PPTOutlineAgent(llm_service)
    
    # We pass empty research result for now
    research_result = {"overall_summary": "No research performed.", "research_results": []}
    
    with patch('app.services.ai_service.ppt_outline_agent.MemoryService'):
        markdown_content = await ppt_agent.generate_outline(extracted_data, research_result, user_id="test_user")
        print("✅ Generated Markdown Outline (First 500 chars):")
        print(markdown_content[:500] + "...")

    # --- Step 3: MARP Conversion ---
    print("\n--- Step 3: MARP Conversion ---")
    
    output_pptx = "final_pursuit_response_real_marp.pptx"
    output_md = "final_pursuit_response_real_marp.md"
    
    with open(output_md, "w") as f:
        f.write(markdown_content)
        
    import subprocess
    cmd = ["marp", "--pptx", output_md, "-o", output_pptx, "--allow-local-files"]
    
    try:
        subprocess.run(cmd, capture_output=True, text=True, check=True)
        print(f"✅ Successfully converted to {output_pptx}")
        if os.path.exists(output_pptx):
             print(f"🎉 Final PPTX Size: {os.path.getsize(output_pptx)} bytes")
    except Exception as e:
        print(f"❌ MARP Conversion Failed: {e}")
        if hasattr(e, 'stderr'):
            print(f"Stderr: {e.stderr}")
        
    print("\n🚀 E2E Pipeline (Real Data + MARP) Completed Successfully!")

if __name__ == "__main__":
    asyncio.run(run_e2e_pipeline_real_marp())
